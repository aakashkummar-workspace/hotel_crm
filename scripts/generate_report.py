"""
Generates a non-technical Word report about the Hotel CRM application.
Output: ./Hotel-CRM-Report.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

GOLD = RGBColor(0xC9, 0xA9, 0x6E)
GOLD_DARK = RGBColor(0x8A, 0x6F, 0x3C)
INK = RGBColor(0x29, 0x26, 0x1B)
INK_LIGHT = RGBColor(0x6E, 0x64, 0x50)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = GOLD_DARK if level <= 2 else INK
        run.font.name = "Georgia"
    return h


def add_para(doc, text, *, bold=False, italic=False, color=None, size=11, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_kv_table(doc, rows, col_widths=(Cm(4.5), Cm(11))):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.width = col_widths[0]
        c1.width = col_widths[1]
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = INK_LIGHT
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.size = Pt(10)
    return table


def add_section_table(doc, rows, head=("Section", "What you can do here")):
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = False
    h = table.rows[0].cells
    h[0].width = Cm(4.5)
    h[1].width = Cm(12)
    h[0].text = head[0]
    h[1].text = head[1]
    for i, (mod, desc) in enumerate(rows, start=1):
        c0, c1 = table.rows[i].cells
        c0.width = Cm(4.5)
        c1.width = Cm(12)
        c0.text = mod
        c1.text = desc
        for run in c0.paragraphs[0].runs:
            run.bold = True
    return table


def add_hr(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9A96E")
    pBdr.append(bottom)
    p_pr.append(pBdr)


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    # ── Title page ──────────────────────────────────────────────────────────
    add_para(doc, "", size=11)
    add_para(doc, "HOTEL", color=GOLD_DARK, size=64, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "CRM", color=GOLD_DARK, size=64, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "", size=11)

    add_para(doc,
             "A complete management system for boutique hotels",
             italic=True, color=INK_LIGHT, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "", size=11)
    add_hr(doc)
    add_para(doc, "", size=11)

    add_para(doc,
             "Run the front desk, the booking calendar, the in-house café, the event hall, "
             "guest profiles, expenses and invoices — all from one app. Plus a beautiful public "
             "booking page guests can use to reserve rooms directly.",
             color=INK_LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "", size=11)
    add_kv_table(doc, [
        ("Application name", "Hotel CRM"),
        ("Designed for", "11-room boutique heritage hotel with café & event hall"),
        ("Who uses it",
         "Hotel managers, front desk staff, café manager, housekeeping supervisor, owners"),
        ("How guests use it",
         "Public booking page — pick dates, choose a room, send an enquiry"),
        ("Languages & currency", "English, Indian Rupees (₹)"),
        ("Where it runs",
         "Any laptop, desktop, tablet or phone with a modern web browser"),
    ], col_widths=(Cm(4.5), Cm(12.5)))

    doc.add_page_break()

    # ── 1. Overview ─────────────────────────────────────────────────────────
    add_heading(doc, "1. What this application does", 1)
    add_para(doc,
        "Hotel CRM is a single application that replaces the spreadsheets, paper diaries and "
        "WhatsApp chats that small hotels normally use to run the property. The whole hotel "
        "lives inside one screen: rooms, today's check-ins, café orders, event bookings, "
        "expenses, invoices, and a guest address book.")
    add_para(doc,
        "It comes in two parts. Staff sign in to the management dashboard. Guests visit a "
        "separate public booking page that shows the rooms, prices and lets them request a stay.")

    add_heading(doc, "Two ways the app is used", 2)
    add_section_table(doc, [
        ("Staff dashboard",
         "Everyone on the team — manager, front desk, café — signs in here with their email "
         "and password. They see today's overview, manage every booking, take café payments, "
         "log expenses, send invoices, look up guest history, and adjust hotel settings."),
        ("Public booking site",
         "A separate page guests open in their browser. Beautiful hero photos, the rooms, the "
         "experiences and reviews, plus a calendar where they pick check-in / check-out dates "
         "and a step-by-step guest counter. Pressing Reserve sends an enquiry directly into the "
         "staff Bookings list as a Pending request."),
    ])

    doc.add_page_break()

    # ── 2. Sections of the staff dashboard ────────────────────────────────
    add_heading(doc, "2. The staff dashboard, section by section", 1)
    add_para(doc,
        "The dashboard has a sidebar with eleven sections. Below is what each one does in "
        "everyday language.")

    # Dashboard
    add_heading(doc, "Dashboard (homepage)", 2)
    add_para(doc, "The first screen after logging in. A live snapshot of the hotel today.")
    for b in [
        "A friendly greeting with today's date",
        "Four big numbers: today's revenue, this month's revenue, total expenses, net profit",
        "Each big number is clickable — opens a pop-up with the breakdown behind it (top days, payment methods, recent transactions and so on)",
        "A revenue chart showing the last 30, 90 days or year-to-date — switch with one click",
        "A circular gauge showing how full the hotel is right now",
        "Counters for available, occupied, reserved and cleaning rooms",
        "Three module cards (Rooms / Coffee Shop / Mini Hall) showing each one's revenue",
        "An Upcoming Check-ins panel — each row has a one-click 'Check in' button that flips the room from 'reserved' to 'occupied' instantly",
        "A Recent Transactions panel from the invoices",
        "Booking sources (which channels brought in the bookings — Direct, Booking.com, etc.)",
        "An Expense breakdown bar chart",
        "An Export button that downloads a CSV of the dashboard data",
    ]:
        add_bullet(doc, b)

    # Rooms
    add_heading(doc, "Rooms", 2)
    add_para(doc, "Everything about the physical rooms — the heart of the hotel.")
    for b in [
        "A row of inventory cards at the top, one per room type (Heritage Suite, Courtyard Deluxe, Garden Twin, Balcony King). Each shows the nightly rate and how many rooms of that type are free right now",
        "Click a type card to edit that type — change the nightly rate, change the bed style, the room size, the maximum guests. Saves instantly and applies to all rooms of that type",
        "Below, every individual room appears as a card with its photo, status, current guest (if any), and price",
        "Filter chips at the top: All, Available, Occupied, Reserved, Cleaning — one click filters the grid",
        "Filter by floor with a drop-down",
        "Three view modes: Grid (photos), List (table), Calendar (a 2-week strip showing who's in which room and when)",
        "Click a room card to open a side panel with the full details: a big photo, rate, beds, size, max guests, amenities and the current guest",
        "Inside the side panel you can:",
        "    – Change the photo: paste an image link, or upload from your device. Pictures are auto-resized so even a 20 MB phone photo works.",
        "    – Move the focal point of the photo: a small target marker lets you pick which part of the photo stays visible when it's cropped to a card",
        "    – Set a custom rate just for this one room (different from the type default). A small 'custom' tag shows when an override is in place",
        "    – Set a custom max-guest count just for this one room (e.g. add a cot)",
        "    – Change the room status with one tap (available / occupied / reserved / cleaning)",
        "    – Press Book this room to jump to the booking form with this room pre-selected",
        "An 'Add room' button opens a small form to add a brand-new room",
        "QR menu and Share booking link buttons copy a public link to your clipboard so you can hand it to guests",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # Bookings
    add_heading(doc, "Bookings", 2)
    add_para(doc, "Every reservation lives here — past, current and upcoming.")
    for b in [
        "A clean table of all bookings: ID, guest name + phone, room, check-in, check-out, nights, amount, source, status",
        "Tabs along the top: All, Checked-in, Confirmed, Pending — each shows a count",
        "Filter by booking source (Direct, Booking.com, WhatsApp, MakeMyTrip etc.)",
        "Tick boxes on each row to multi-select for bulk export",
        "Export button downloads the visible bookings (or the selected ones) as a spreadsheet",
        "A '…' menu on each row lets you: mark as checked-in, mark as checked-out, mark as confirmed/pending, message the guest on WhatsApp directly, or delete the booking",
        "Deleting prompts a confirmation dialog so nothing is removed by accident",
        "New booking button opens a form:",
        "    – Guest name, phone, email",
        "    – Pick check-in and check-out dates with a date picker",
        "    – The system then automatically shows only the rooms that are genuinely free for those dates — sold-out types are flagged",
        "    – Pick the source and number of guests",
        "    – The total cost (room rate × nights + tax) is calculated as you type",
        "When you save a booking, the system automatically:",
        "    – Marks the room as 'reserved' with the guest's name and dates",
        "    – Adds the guest to the Guests CRM (or increments their visits + lifetime spend if they're a returning guest)",
        "    – Records the action in the Activity log",
        "If you ever try to book a room that's already reserved for overlapping dates, the system blocks it with a clear message",
    ]:
        add_bullet(doc, b)

    # Coffee Shop
    add_heading(doc, "Coffee Shop", 2)
    add_para(doc, "A point-of-sale screen for the in-house café — looks and works like a real café till.")
    for b in [
        "Live numbers at the top: orders today, revenue today, average order value, top item",
        "Category chips: All, Espresso, Filter, Tea, Cold, Pastry, All-Day",
        "Menu items appear as cards — tap to add to the cart on the right",
        "Cart shows quantity steppers, line totals, GST (5%), and a final total",
        "Pick a table number (T1–T5, Takeaway, or charge to a hotel room)",
        "Optional customer name field",
        "Pay with UPI, Card, Cash, or charge to a hotel room",
        "Charge button completes the order and clears the cart",
        "Z-Report button opens a daily summary: total orders, total revenue, breakdown by payment method — and can be exported as CSV",
        "Menu button shows the entire café menu in one place",
    ]:
        add_bullet(doc, b)

    # Mini Hall
    add_heading(doc, "Mini Hall", 2)
    add_para(doc, "Reservations for the property's event hall (weddings, birthdays, corporate offsites).")
    for b in [
        "A hero card at the top with the hall's photo, capacity, and rate card (half day, full day, evening, wedding package)",
        "Live totals: events booked this month, revenue, advance collected",
        "A list of all upcoming events with title, contact person, date, time, guest count, and a progress bar showing how much of the total has been paid as advance",
        "Click any event row to open a side drawer with full details and a one-click Confirm or Mark Pending button",
        "Reserve hall button opens a form to book a new event",
        "Hall booking link button copies a shareable URL so couples can browse and enquire from their phone",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # Expenses
    add_heading(doc, "Expenses", 2)
    add_para(doc, "Tracks every rupee leaving the hotel — salaries, utilities, supplies, repairs.")
    for b in [
        "Big-number tiles: this month's spend, the largest category, daily average, pending bills",
        "A by-category bar chart and a profit waterfall (revenue in − expenses out = net profit)",
        "Live category filter — pick Salaries, Utilities, Cleaning, Maintenance, Coffee Purchases or Misc",
        "Add expense button: vendor, date, category, amount, payment method, optional note",
        "A '…' menu on each row to delete (with confirmation)",
        "Export button saves a CSV of the visible expenses",
    ]:
        add_bullet(doc, b)

    # Guests CRM
    add_heading(doc, "Guests CRM", 2)
    add_para(doc, "An address book and history file for every person who has ever stayed at the hotel.")
    for b in [
        "Searchable list on the left, full guest profile on the right",
        "Header banner with the guest's name, contact details, and a tier badge (VIP / Regular / New)",
        "Quick-action buttons — WhatsApp, Email, Call — open the right app on your device",
        "Lifetime spend, visit count, average booking value, last stay date",
        "Concierge note for important context (allergies, anniversaries, preferences)",
        "Communication history showing recent WhatsApp / email / call interactions",
        "New booking shortcut that pre-fills the booking form with this guest's details",
        "Add guest button to manually create a profile",
        "Send campaign button opens an email composed to every guest on file (BCC)",
    ]:
        add_bullet(doc, b)

    # Reports
    add_heading(doc, "Reports", 2)
    add_para(doc, "Charts and analytics, the kind of thing owners and accountants ask for.")
    for b in [
        "Period switcher: 7 days, 30 days, 90 days, year-to-date — every chart updates",
        "Total revenue, total expenses, net profit, average occupancy",
        "Stacked revenue chart by module (rooms / café / hall)",
        "Donut showing the module mix",
        "Occupancy bars over the last 12 weeks",
        "Expense breakdown bars",
        "PDF button — opens the print dialog so you can save as PDF",
        "CSV button — downloads the chart data",
        "Six saved-report cards (Monthly P&L, Occupancy, Café performance, Guest cohorts, Expense audit, Hall pipeline). Click one to download its data instantly",
    ]:
        add_bullet(doc, b)

    # Invoices
    add_heading(doc, "Invoices & Payments", 2)
    add_para(doc, "Receipts and bills the hotel issues to guests.")
    for b in [
        "Big-number tiles: total billed, collected, pending, average invoice",
        "Status filter chips: All, Paid, Pending, Partial, Advance",
        "A neat table showing each invoice's number, guest, date, subtotal, tax, total, payment method, status",
        "Click any row to open a printable invoice view with the hotel logo, GST number, line items, totals",
        "Mark as paid button changes the status with one click",
        "Download PDF opens the print dialog so you can save the invoice as a PDF",
        "Send button records that the invoice has been sent",
        "New invoice form: guest, date, amount, tax (auto-calculated at 18% if left blank), method, status",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # Settings
    add_heading(doc, "Settings", 2)
    add_para(doc, "Hotel-wide preferences. Six tabs.")
    for b in [
        "Hotel profile — name, location, tagline, contact email, phone, GST number, currency, an 'about' description. The profile shows on invoices and the public booking page",
        "Team & roles — see all staff accounts at a glance",
        "Integrations — connect / disconnect WhatsApp Business, Razorpay, Booking.com, Google Calendar, Cloudinary, Mailchimp",
        "Billing & tax — room GST %, café GST %, hall GST %, invoice number prefix and next number",
        "Security — toggles for two-factor authentication, idle session timeout, IP allow-list, encryption at rest",
        "Activity log — a chronological audit trail of every important action (bookings created, status changes, deletes, invoices sent, etc.) with the staff member who did it",
    ]:
        add_bullet(doc, b)

    # ── 3. Public booking site ────────────────────────────────────────────
    add_heading(doc, "3. The public booking site", 1)
    add_para(doc,
        "A separate, beautiful page guests can visit directly — no login required. Designed to "
        "look like a premium hotel website, not a CRM screen.")
    for b in [
        "Full-width hero with a large photo, the property name, and a five-star reviewer badge",
        "Three quick stats: number of restored rooms, average review score, walking distance to the sea",
        "A floating booking widget pinned across the hero: Check-in, Check-out, Guests, Rooms, Search",
        "Date picker shows two calendar months side by side and lets you drag a date range across them",
        "Guest stepper — separate counters for Adults, Children, Infants",
        "Rooms section showing every room type as a card with photo, key features, description, price and a Reserve button",
        "Experiences section (sunrise walk, filter-coffee craft, old-town tour, ayurvedic supper, headland bicycle ride)",
        "Our Story section explaining the heritage and history of the building",
        "Reviews section with starred guest quotes",
        "Footer with directions, WhatsApp link, links to Stay / House / Help sections",
        "When a guest presses Reserve, an inline form appears: name, phone, email, optional notes. After submit they see a confirmation screen with a reference number — and the enquiry shows up immediately in the staff Bookings list as a Pending booking",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # ── 4. Booking lifecycle ────────────────────────────────────────────
    add_heading(doc, "4. How a booking moves through the system", 1)
    add_para(doc,
        "Bookings have a status that moves through four stages. The system updates room status "
        "and the guest record automatically — staff don't need to remember to do anything else.")

    add_section_table(doc, [
        ("1. Pending",
         "Guest submits an enquiry from the public site, or staff create a tentative booking. "
         "Room shows as 'Reserved' to keep it from being double-booked."),
        ("2. Confirmed",
         "Staff confirm the booking (after taking deposit, calling back, etc.). Room stays "
         "'Reserved'. Guest visit count goes up; their lifetime spend increases."),
        ("3. Checked-in",
         "Guest arrives. One click on the dashboard or in the booking row flips the booking "
         "to 'checked-in' and the room to 'Occupied'. The guest's name and check-out date appear "
         "on the room card."),
        ("4. Checked-out",
         "Guest leaves. Room flips to 'Cleaning' so housekeeping knows it needs to be turned "
         "over. Once housekeeping marks it 'Available', it becomes bookable again."),
    ], head=("Stage", "What happens"))

    add_para(doc, "")
    add_para(doc,
        "Cancellations free the room immediately. If a booking is deleted, the system frees "
        "the room. Trying to book the same room for overlapping dates as an existing booking "
        "is blocked with a clear message naming the conflicting reservation.")

    # ── 5. Customisations ───────────────────────────────────────────────────
    add_heading(doc, "5. Customisation owners can do without help", 1)
    add_para(doc, "Almost everything in the app can be edited by the owner from the screen — no developer needed.")
    for b in [
        "Hotel name, tagline, location, contact details, GST number, about-text — all editable in Settings",
        "Tax rates for rooms, café and hall — separate percentages, all editable",
        "Invoice number prefix and the next number to use — editable",
        "Room types: name, default rate, default size, default bed, default max guests",
        "Per-room overrides: a different rate for one specific room, a different max-guest count",
        "Room photos — upload from the device or paste a link. Auto-resized so any size file works",
        "Room photo focal point — drag a marker over the photo to pick which part stays visible when cropped to a card",
        "Add new rooms or remove old ones",
        "Café menu items, prices and categories",
        "Switch between Dark and Light theme",
        "Three sidebar styles (full-width, icon-only, floating)",
        "Four font pairings (typeface combinations) for the typography",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # ── 6. Search & navigation ────────────────────────────────────────────
    add_heading(doc, "6. Search and quick navigation", 1)
    for b in [
        "A search box in the top bar (and in the sidebar) opens a global command palette",
        "Press Ctrl+K (or Cmd+K) anywhere to open it",
        "Type and the system instantly searches across rooms, bookings, guests, invoices",
        "Or just type the name of a section ('rooms', 'invoices') to jump straight there",
        "Use arrow keys to navigate, Enter to open, Esc to close",
    ]:
        add_bullet(doc, b)

    # ── 7. Notifications & activity ───────────────────────────────────────
    add_heading(doc, "7. Notifications & activity log", 1)
    for b in [
        "A bell icon in the top bar shows unread notifications: new bookings, payments received, check-ins due, low stock, recent reviews",
        "A 'Mark all read' button clears the unread badges",
        "The Activity Log inside Settings keeps a permanent record of every booking creation, status change, deletion, invoice sent, and rate change — with the staff member's name and timestamp",
    ]:
        add_bullet(doc, b)

    # ── 8. Look & feel ─────────────────────────────────────────────────────
    add_heading(doc, "8. Look and feel", 1)
    add_para(doc,
        "The application is designed to look like a high-end heritage stay rather than a generic admin tool. "
        "Warm gold accents, serif headings, soft cream backgrounds in light mode and warm dark tones at night.")
    for b in [
        "Dark mode and light mode — switch in one click from the top bar",
        "Three sidebar layouts: full (with labels), collapsed (icons only), or floating (a card that hovers over the page)",
        "Four typography pairings — pick a different combination of display & body fonts",
        "Animations on page changes and hover effects on every interactive element",
        "Looks great on phones, tablets, laptops, and big monitors",
    ]:
        add_bullet(doc, b)

    # ── 9. Exports ─────────────────────────────────────────────────────────
    add_heading(doc, "9. What you can export", 1)
    add_para(doc, "Every list-style screen has an Export button. You also get printable views for invoices and reports.")
    for b in [
        "Dashboard — daily revenue table for the chosen period",
        "Bookings — visible bookings (or the selected rows) as a spreadsheet",
        "Guests — searchable list (via the saved Guest cohorts report)",
        "Coffee Shop — Z-Report (orders + payment-method breakdown)",
        "Expenses — visible expenses, filtered by category",
        "Invoices — all invoices, plus an individual printable invoice with QR code area",
        "Reports — six saved reports each with their own export, plus a 'Save as PDF' for the entire reports page",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # ── 10. Hotel policies & guest privileges ─────────────────────────────
    add_heading(doc, "10. Hotel policies & guest privileges", 1)
    add_para(doc,
        "These are the house policies the hotel runs by. The application is designed to "
        "support each one — staff don't have to remember edge cases by heart, the system "
        "flags them at the right moment.")

    add_heading(doc, "Live booking status visible to guests", 2)
    for b in [
        "Every guest can see the live status of their reservation at any time: Pending, Confirmed, Checked-in or Checked-out",
        "The status updates automatically as the front desk processes the booking — guests don't need to call to ask",
        "On the public site, returning guests with a booking reference can pull up their reservation card and see exactly where it stands",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "Booking history with filters", 2)
    for b in [
        "Returning guests can view their complete history of stays at the hotel",
        "Filter by date range (last month, last six months, last year, custom)",
        "Filter by status — upcoming, completed, cancelled",
        "Filter by room type — Heritage Suite, Courtyard Deluxe, etc.",
        "Each entry shows the dates, room number, nights, total paid, payment method, and any concierge notes",
        "Staff see the same history inside the Guests CRM section — useful for personalising returning visits",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "Two-hour flexible check-out (free)", 2)
    for b in [
        "When a guest's stay ends, the first two hours past the standard check-out time are granted free of charge",
        "Lets guests breathe — finish breakfast, take a last walk, wait for an evening flight",
        "After the two-hour grace window the system starts charging — the room rate is pro-rated to the additional hours used",
        "The dashboard flags late check-outs so the front desk knows to add the late fee at settlement",
        "The hotel sets the late-check-out rate inside Settings → Billing & tax, so the policy can be tuned for peak vs. off-peak season",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "Installment billing for long stays (15 nights or more)", 2)
    for b in [
        "Any booking of 15 or more consecutive nights automatically qualifies for split payment",
        "Half of the total is collected as an advance at the time of booking",
        "The remaining half is collected at the time of check-out",
        "The system creates two invoice rows automatically — the first stamped 'Advance', the second stamped 'Balance due'",
        "When the advance is recorded, the booking moves from Pending to Confirmed",
        "When the balance is paid at check-out, the second invoice is marked Paid and the booking is closed",
        "Owners can see at a glance which long stays still owe a balance from the Invoices screen",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "Cloud-kitchen add-on (activated when the service starts)", 2)
    for b in [
        "The application is ready for an in-house cloud kitchen module",
        "When the cloud-kitchen business starts operating, a dedicated kitchen account is created with its own login",
        "Kitchen staff sign in to a focused screen — incoming orders, prep queue, dispatch board",
        "Orders flow through the same point-of-sale as the café, so payments and tax handling are consistent",
        "Kitchen inventory, ingredient costs and revenue automatically roll into the existing Expenses and Reports modules",
        "No retraining for management — the dashboard simply shows another module card alongside Rooms, Coffee Shop and Mini Hall",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "Complimentary vehicle for long-stay guests", 2)
    for b in [
        "Guests staying 15 nights or more can request a hotel-provided car for short trips, airport transfers, or sightseeing",
        "The request is captured against the booking and shows up on the front desk's daily checklist",
        "Driver assignment, departure time, return time and trip purpose are logged under the booking notes",
        "A fuel and mileage log is kept against each car so the hotel can track running costs from the Expenses screen",
        "If multiple long-stay guests request the car at overlapping times, the system flags the conflict so staff can coordinate",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # ── 11. Future plans ──────────────────────────────────────────────────
    add_heading(doc, "11. Things planned for later versions", 1)
    for b in [
        "Drag-and-drop bookings on the room calendar to extend or move stays",
        "Real Booking.com / MakeMyTrip channel sync (rates, availability, bookings flow both ways)",
        "Take payments online: Razorpay deposit collection on the public booking page",
        "Auto-send WhatsApp messages: confirmation when booking is made, reminder day before, thank-you after stay",
        "Multi-property support so one owner can run several hotels from the same dashboard",
        "Different access levels for different staff roles (admin / front desk / housekeeping / café)",
        "A live notifications feed on the dashboard whenever a new booking comes in from the public site",
        "Pre-built guest cohort and retention analytics",
        "Online room photo storage on a separate image service for very large galleries",
    ]:
        add_bullet(doc, b)

    # ── Closing ────────────────────────────────────────────────────────────
    add_heading(doc, "In summary", 1)
    add_para(doc,
        "Hotel CRM gives a small heritage hotel everything it needs to run day to day from a "
        "single warm, easy-to-use screen. Every button does something real — bookings update "
        "rooms, rooms update the guest record, the public site feeds straight into the staff "
        "list. Owners can change rates, photos, tax and almost anything else themselves "
        "without calling for help.")
    add_para(doc,
        "It's been built so that adding new modules later (a spa, a restaurant, a gift shop) "
        "follows the same pattern as everything that's already there.",
        italic=True, color=INK_LIGHT)

    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    out_path = os.path.join(base_dir, "Hotel-CRM-Report.docx")
    try:
        doc.save(out_path)
        print(f"OK: {out_path}")
    except PermissionError:
        # File is open in Word — save next to it under a different name so the user
        # still gets the new version without having to close Word first.
        fallback = os.path.join(base_dir, "Hotel-CRM-Report-new.docx")
        doc.save(fallback)
        print(f"OK (original locked, wrote alternate): {fallback}")


if __name__ == "__main__":
    build()
